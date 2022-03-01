import re

from docx import Document
from docxtpl import DocxTemplate, InlineImage

from core.settings import DOCUMENT_ROOT, TEMPLATE_ROOT, MEDIA_ROOT
from objects.record.get_record import get_object_record_by_id_http, get_record_title
from data_base_driver.sys_key.get_key_dump import get_key_by_id
from objects.relations.get_rel import get_object_relation


def get_document_from_template(template_name, title, data):
    """
    Функция для формирования документа microsoft word из шаблона
    @param template_name: название шаблона
    @param title: название документа
    @param data: словарь содержащий информаию для занесение в документ, формат:
    {Position:[data1,...,dataN],...,PositionN:p[]} , где Position позиция в шаблоне, data -  данные для занесения
    @return: путь к конечному документу
    """
    document = Document(TEMPLATE_ROOT + template_name)
    for paragraph in document.paragraphs:
        for word in paragraph.text.replace('\t', ' ').split(' '):
            if data.get(re.sub("{[^<]+}", "", word)):
                size = paragraph.runs[0].font.size
                font_name = paragraph.style.font.name
                try:
                    paragraph.text = paragraph.text.replace(word, str(data[word][0]) if word.find('{') == -1 else
                    str(data[re.sub("{[^<]+}", "", word)][int(word[word.find('{') + 1:word.find('}')]) - 1]))
                except IndexError as e:
                    raise Exception(1, 'Выход за пределы массива')
                paragraph.style.font.name = font_name
                for run in paragraph.runs:
                    run.font.size = size
    document.save(DOCUMENT_ROOT + title + '.docx')
    return DOCUMENT_ROOT + title + '.docx'


def get_dossier_for_object(group_id, object_id, rec_id, title):
    """
    Функция для формирования досье о любом объекте базы данных
    @param group_id: идентификатор группы пользователя
    @param object_id: идентификатор типа объекта
    @param rec_id: идентификатор объекта
    @param title: название конечного досье
    @return: путь к конечному файлу
    """
    template = DocxTemplate(TEMPLATE_ROOT + 'template_tables.docx')
    object_records = get_object_record_by_id_http(object_id, rec_id, group_id)
    table_content = []
    for param in object_records['params']:
        if get_key_by_id(param['id'])['type'] == 'file_photo':
            value = InlineImage(template,
                                MEDIA_ROOT + '/files/' + str(object_id) + '/' + str(rec_id) + '/' + param['values'][0][
                                    'value'])
        elif get_key_by_id(param['id'])['type'] == 'file_any':
            continue
        else:
            value = param['values'][0]['value']
        table_content.append({'id': param['id'], 'key': get_key_by_id(param['id'])['title'], 'value': value})
    object_relations = get_object_relation(group_id, object_id, rec_id, [], True)
    object_relations.sort(key=lambda x: x['object_id'])
    relations = []
    for related_object in object_relations:
        for relation in related_object['relations']:
            title_object = get_record_title(related_object['object_id'], related_object['rec_id'], group_id)['title']
            key_title = get_key_by_id(relation['id'])['title'] + ' ' + relation['values'][0]['value'] + ' ' + \
                        relation['values'][0]['date']
            relations.append({'id': relation['id'], 'key': key_title, 'value': title_object})
    table_content.sort(key=lambda x: get_key_by_id(x['id'])['type'] != 'file_photo')
    content = {
        'title': 'Досье на ' + object_records['title'],
        'table_contents': table_content,
        'relations': relations
    }
    template.render(content)
    template.save(DOCUMENT_ROOT + title + '.docx')
    return DOCUMENT_ROOT + title + '.docx'


def get_docx_document_text(path):
    result = ''
    document = Document(path)
    for paragraph in document.paragraphs:
        result += paragraph.text + '\n'
    return result
